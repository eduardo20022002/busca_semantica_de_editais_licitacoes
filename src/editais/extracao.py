from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath


@dataclass(frozen=True)
class TextoExtraido:
    nome_arquivo: str
    texto: str


def extrair_textos_de_arquivo(nome_arquivo: str, conteudo: bytes) -> list[TextoExtraido]:
    extensao = PurePosixPath(nome_arquivo).suffix.lower()

    if extensao == ".zip":
        return _extrair_de_zip(conteudo)

    extrator = _EXTRATORES.get(extensao)
    if extrator is None:
        return []

    try:
        texto = extrator(conteudo)
    except Exception:
        # Arquivo corrompido, protegido por senha, ou variante não suportada do
        # formato: pula sem derrubar o edital (política de "pular e registrar" da
        # Análise profunda — sem OCR nesta versão).
        return []

    if not texto.strip():
        return []
    return [TextoExtraido(nome_arquivo=nome_arquivo, texto=texto)]


def _extrair_txt(conteudo: bytes) -> str:
    return conteudo.decode("utf-8", errors="ignore")


def _extrair_pdf(conteudo: bytes) -> str:
    from pypdf import PdfReader

    leitor = PdfReader(io.BytesIO(conteudo))
    return "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)


def _extrair_docx(conteudo: bytes) -> str:
    from docx import Document

    documento = Document(io.BytesIO(conteudo))
    partes = [paragrafo.text for paragrafo in documento.paragraphs]
    for tabela in documento.tables:
        for linha in tabela.rows:
            partes.extend(celula.text for celula in linha.cells)
    return "\n".join(partes)


def _extrair_xlsx(conteudo: bytes) -> str:
    from openpyxl import load_workbook

    livro = load_workbook(io.BytesIO(conteudo), read_only=True, data_only=True)
    partes: list[str] = []
    for planilha in livro.worksheets:
        for linha in planilha.iter_rows(values_only=True):
            partes.extend(str(celula) for celula in linha if celula is not None)
    livro.close()
    return "\n".join(partes)


_EXTRATORES = {
    ".txt": _extrair_txt,
    ".pdf": _extrair_pdf,
    ".docx": _extrair_docx,
    ".xlsx": _extrair_xlsx,
}


def _extrair_de_zip(conteudo: bytes) -> list[TextoExtraido]:
    try:
        zip_file = zipfile.ZipFile(io.BytesIO(conteudo))
    except zipfile.BadZipFile:
        return []

    textos: list[TextoExtraido] = []
    with zip_file:
        for nome_interno in zip_file.namelist():
            if nome_interno.endswith("/"):
                continue
            try:
                bytes_internos = zip_file.read(nome_interno)
            except Exception:
                continue
            textos.extend(extrair_textos_de_arquivo(nome_interno, bytes_internos))
    return textos
